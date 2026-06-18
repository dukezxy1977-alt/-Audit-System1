# 文件处理模块 - 支持 PDF、Word、TXT 文件

import os
import zipfile

import pdfplumber
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class FileProcessor:
    """处理各种文件格式"""

    @staticmethod
    def extract_text(file_path):
        """
        从文件中提取文本及分段位置信息。
        支持：PDF, DOCX, DOC, TXT

        返回 (full_text, segments)：
          - full_text: 拼接后的全文，始终是字符串（不会是 None）
          - segments: [{'location': '第3页' / '第5段' / '第2行', 'text': '...'}, ...]
            用于在发现风险时定位“在文档的什么位置”
        """

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.pdf':
            segments = FileProcessor._extract_pdf(file_path)
        elif file_ext == '.docx':
            segments = FileProcessor._extract_word(file_path)
        elif file_ext == '.doc':
            # python-docx 无法解析旧版二进制 .doc 格式，明确提示用户而不是抛出底层异常
            raise ValueError(
                "不支持旧版 .doc 格式（仅支持 Word 2007+ 的 .docx）。"
                "请使用 Word 将文件“另存为” .docx 格式后重新上传"
            )
        elif file_ext == '.txt':
            segments = FileProcessor._extract_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext or '(无扩展名)'}")

        full_text = "\n".join(seg['text'] for seg in segments)

        if not full_text.strip():
            raise ValueError(
                "未能从文件中提取到任何文本内容，"
                "该文件可能是扫描版/图片版文档，或文件已损坏"
            )

        return full_text, segments

    @staticmethod
    def _extract_pdf(file_path):
        """从 PDF 文件提取文本，按页分段"""
        try:
            segments = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    # pdfplumber 对于纯图片/扫描页会返回 None，必须做空值兜底
                    page_text = page.extract_text() or ""
                    segments.append({'location': f'第{i}页', 'text': page_text})
            return segments
        except Exception as e:
            raise Exception(f"PDF 提取失败: {str(e)}") from e

    @staticmethod
    def _extract_word(file_path):
        """从 Word (.docx) 文件提取文本，按段落/表格单元格分段"""
        try:
            doc = Document(file_path)

            segments = []
            for i, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    segments.append({'location': f'第{i}段', 'text': para.text})

            # 补充提取表格中的文本，避免表格内容被遗漏
            for t, table in enumerate(doc.tables, 1):
                for r, row in enumerate(table.rows, 1):
                    for c, cell in enumerate(row.cells, 1):
                        if cell.text.strip():
                            segments.append({
                                'location': f'表格{t}第{r}行第{c}列',
                                'text': cell.text
                            })

            return segments
        except (PackageNotFoundError, zipfile.BadZipFile) as e:
            raise Exception(
                "Word 提取失败: 文件不是有效的 .docx 格式"
                "（可能是旧版 .doc 文件被重命名为 .docx，或文件已损坏）"
            ) from e
        except Exception as e:
            raise Exception(f"Word 提取失败: {str(e)}") from e

    @staticmethod
    def _extract_txt(file_path):
        """从 TXT 文件读取文本，自动兼容 UTF-8 / GBK 等常见中文编码，按行分段"""
        last_error = None
        for encoding in ('utf-8', 'utf-8-sig', 'gb18030'):
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    lines = f.readlines()
                return [
                    {'location': f'第{i}行', 'text': line}
                    for i, line in enumerate(lines, 1)
                    if line.strip()
                ]
            except UnicodeDecodeError as e:
                last_error = e
                continue
            except Exception as e:
                raise Exception(f"TXT 读取失败: {str(e)}") from e
        raise Exception(f"TXT 读取失败: 无法识别文件编码 ({last_error})")
